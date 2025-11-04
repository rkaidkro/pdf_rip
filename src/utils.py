"""
Utility functions for the PDF processing pipeline.
"""

import hashlib
import json
import logging
import os
import time
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import psutil
from loguru import logger
import fitz  # PyMuPDF


def setup_logging(level: str = "INFO", log_file: Optional[Path] = None) -> None:
    """Setup logging configuration with rich formatting."""
    # Remove default handler
    logger.remove()
    
    # Add console handler with rich formatting
    logger.add(
        lambda msg: print(msg, end=""),
        level=level,
        format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <cyan>{name}</cyan>:<cyan>{function}</cyan>:<cyan>{line}</cyan> - <level>{message}</level>",
        colorize=True
    )
    
    # Add file handler if specified
    if log_file:
        logger.add(
            log_file,
            level=level,
            format="{time:YYYY-MM-DD HH:mm:ss} | {level: <8} | {name}:{function}:{line} - {message}",
            rotation="10 MB",
            retention="7 days"
        )


def generate_run_id() -> str:
    """Generate a unique run identifier."""
    return f"run_{uuid.uuid4().hex[:8]}_{int(time.time())}"


def calculate_element_hash(content: str, bbox: List[float], page: int) -> str:
    """Calculate a hash for an element based on content and position."""
    hash_input = f"{content}:{bbox}:{page}".encode('utf-8')
    return hashlib.sha256(hash_input).hexdigest()[:16]


def get_memory_usage() -> float:
    """Get current memory usage in MB."""
    process = psutil.Process(os.getpid())
    return process.memory_info().rss / 1024 / 1024


def validate_pdf_file(pdf_path: Path) -> Tuple[bool, str]:
    """Validate that a PDF file exists and is readable."""
    if not pdf_path.exists():
        return False, f"PDF file does not exist: {pdf_path}"
    
    if not pdf_path.is_file():
        return False, f"Path is not a file: {pdf_path}"
    
    if pdf_path.stat().st_size == 0:
        return False, f"PDF file is empty: {pdf_path}"
    
    try:
        # Try to open with PyMuPDF to validate
        doc = fitz.open(str(pdf_path))
        page_count = len(doc)
        doc.close()
        
        if page_count == 0:
            return False, f"PDF has no pages: {pdf_path}"
        
        return True, f"Valid PDF with {page_count} pages"
    except Exception as e:
        return False, f"Failed to open PDF: {str(e)}"


def validate_word_file(docx_path: Path) -> Tuple[bool, str]:
    """Validate that a Word document exists and is readable."""
    if not docx_path.exists():
        return False, f"Word document does not exist: {docx_path}"
    
    if not docx_path.is_file():
        return False, f"Path is not a file: {docx_path}"
    
    if docx_path.stat().st_size == 0:
        return False, f"Word document is empty: {docx_path}"
    
    try:
        # Try to open with python-docx to validate
        import docx
        doc = docx.Document(str(docx_path))
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        
        if paragraph_count == 0 and table_count == 0:
            return False, f"Word document has no content: {docx_path}"
        
        return True, f"Valid Word document with {paragraph_count} paragraphs and {table_count} tables"
    except Exception as e:
        return False, f"Failed to open Word document: {str(e)}"


def validate_document_file(file_path: Path) -> Tuple[bool, str]:
    """Validate that a document file exists and is readable (PDF or Word)."""
    if not file_path.exists():
        return False, f"Document file does not exist: {file_path}"
    
    if not file_path.is_file():
        return False, f"Path is not a file: {file_path}"
    
    if file_path.stat().st_size == 0:
        return False, f"Document file is empty: {file_path}"
    
    # Check file extension
    suffix = file_path.suffix.lower()
    
    if suffix == '.pdf':
        return validate_pdf_file(file_path)
    elif suffix in ['.docx', '.doc']:
        return validate_word_file(file_path)
    else:
        return False, f"Unsupported file format: {suffix}. Supported formats: .pdf, .docx, .doc"


def detect_document_characteristics(pdf_path: Path) -> Dict[str, Any]:
    """Analyze PDF to detect document characteristics."""
    characteristics = {
        "is_scanned": False,
        "has_text_layer": False,
        "page_count": 0,
        "table_density": 0.0,
        "math_signals": 0,
        "languages": [],
        "estimated_domain": None
    }
    
    try:
        doc = fitz.open(str(pdf_path))
        characteristics["page_count"] = len(doc)
        
        text_content = ""
        table_count = 0
        math_count = 0
        
        # Sample first few pages for analysis
        sample_pages = min(3, len(doc))
        
        for page_num in range(sample_pages):
            page = doc[page_num]
            
            # Check for text layer
            text = page.get_text()
            text_content += text
            
            # Count tables (simple heuristic based on table-like structures)
            try:
                tables = page.get_tables()
                table_count += len(tables)
            except AttributeError:
                # PyMuPDF version doesn't support get_tables, skip table counting
                pass
            
            # Look for math symbols
            math_symbols = ['∑', '∫', '∏', '√', '∞', '±', '≤', '≥', '≠', '≈', '→', '←', '↔']
            for symbol in math_symbols:
                if symbol in text:
                    math_count += 1
        
        # Calculate average text per page
        avg_text_per_page = len(text_content.strip()) / sample_pages if sample_pages > 0 else 0
        
        # Consider it scanned if:
        # 1. No text at all, OR
        # 2. Very little text (less than 50 chars per page on average) - likely just page numbers or metadata
        characteristics["has_text_layer"] = len(text_content.strip()) > 0
        characteristics["is_scanned"] = not characteristics["has_text_layer"] or avg_text_per_page < 50
        characteristics["table_density"] = table_count / sample_pages if sample_pages > 0 else 0
        characteristics["math_signals"] = math_count
        
        # Simple language detection (basic)
        if any(char.isascii() for char in text_content[:1000]):
            characteristics["languages"].append("en")
        
        # Domain estimation based on content
        if math_count > 5:
            characteristics["estimated_domain"] = "academic"
        elif table_count > 2:
            characteristics["estimated_domain"] = "business"
        else:
            characteristics["estimated_domain"] = "general"
        
        doc.close()
        
    except Exception as e:
        logger.warning(f"Error analyzing PDF characteristics: {e}")
    
    return characteristics


def ensure_directory(path: Path) -> None:
    """Ensure a directory exists, creating it if necessary."""
    path.mkdir(parents=True, exist_ok=True)


def save_jsonl(data: List[Dict[str, Any]], filepath: Path) -> None:
    """Save data as JSONL (JSON Lines) format."""
    ensure_directory(filepath.parent)
    with open(filepath, 'w', encoding='utf-8') as f:
        for item in data:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')


def load_jsonl(filepath: Path) -> List[Dict[str, Any]]:
    """Load data from JSONL format."""
    data = []
    if filepath.exists():
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    data.append(json.loads(line))
    return data


def calculate_cer(reference: str, hypothesis: str) -> float:
    """Calculate Character Error Rate."""
    if not reference:
        return 1.0 if hypothesis else 0.0
    
    # Simple Levenshtein distance for CER
    m, n = len(reference), len(hypothesis)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    
    for i in range(m + 1):
        dp[i][0] = i
    for j in range(n + 1):
        dp[0][j] = j
    
    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if reference[i - 1] == hypothesis[j - 1]:
                dp[i][j] = dp[i - 1][j - 1]
            else:
                dp[i][j] = 1 + min(dp[i - 1][j], dp[i][j - 1], dp[i - 1][j - 1])
    
    return dp[m][n] / len(reference)


def calculate_wer(reference: str, hypothesis: str) -> float:
    """Calculate Word Error Rate."""
    ref_words = reference.split()
    hyp_words = hypothesis.split()
    
    if not ref_words:
        return 1.0 if hyp_words else 0.0
    
    # Simple word-level Levenshtein
    m, n = len(ref_words), len(hyp_words)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    
    for i in range(m + 1):
        dp[i][0] = i
    for j in range(n + 1):
        dp[0][j] = j
    
    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if ref_words[i - 1] == hyp_words[j - 1]:
                dp[i][j] = dp[i - 1][j - 1]
            else:
                dp[i][j] = 1 + min(dp[i - 1][j], dp[i][j - 1], dp[i - 1][j - 1])
    
    return dp[m][n] / len(ref_words)


def sanitize_filename(filename: str) -> str:
    """Sanitize a filename for safe filesystem use."""
    # Remove or replace problematic characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    # Limit length
    if len(filename) > 200:
        name, ext = os.path.splitext(filename)
        filename = name[:200-len(ext)] + ext
    
    return filename


class Timer:
    """Context manager for timing operations."""
    
    def __init__(self, name: str = "operation"):
        self.name = name
        self.start_time = None
        self.end_time = None
    
    def __enter__(self):
        self.start_time = time.time()
        logger.info(f"Starting {self.name}")
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.end_time = time.time()
        duration = self.end_time - self.start_time
        logger.info(f"Completed {self.name} in {duration:.2f}s")
    
    @property
    def duration(self) -> float:
        """Get the duration in seconds."""
        if self.start_time is None:
            return 0.0
        end = self.end_time or time.time()
        return end - self.start_time
