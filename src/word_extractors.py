"""
Content extractors for Word documents.
"""

import docx
import docx2txt
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from loguru import logger

from .models import ProvenanceRecord
from .utils import calculate_element_hash


class WordTextExtractor:
    """Extract text content from Word documents."""
    
    def __init__(self):
        self.version = "1.0.0"
    
    def get_version(self) -> str:
        return self.version
    
    def extract_text(self, docx_path: Path) -> Dict[str, Any]:
        """Extract text from Word document using python-docx."""
        result = {
            "content": "",
            "provenance": []
        }
        
        try:
            doc = docx.Document(str(docx_path))
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    result["content"] += text + "\n"
                    
                    # Create provenance record for paragraph
                    provenance = ProvenanceRecord(
                        page=1,  # Word docs don't have explicit pages like PDFs
                        bbox=[0, 0, 0, 0],  # Placeholder for Word docs
                        tool="python-docx",
                        confidence=1.0,
                        element_hash=calculate_element_hash(text, [0, 0, 0, 0], 1),
                        element_type="paragraph",
                        content_preview=text[:50]
                    )
                    result["provenance"].append(provenance)
            
            # Extract text from tables
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    table_content = "\n".join(table_text)
                    result["content"] += f"\n[Table {table_num + 1}]\n{table_content}\n"
                    
                    # Create provenance record for table
                    provenance = ProvenanceRecord(
                        page=1,
                        bbox=[0, 0, 0, 0],
                        tool="python-docx",
                        confidence=1.0,
                        element_hash=calculate_element_hash(table_content, [0, 0, 0, 0], 1),
                        element_type="table",
                        content_preview=table_content[:50]
                    )
                    result["provenance"].append(provenance)
            
        except Exception as e:
            logger.error(f"Word text extraction failed: {e}")
            raise
        
        return result
    
    def extract_text_from_bytes(self, docx_bytes: bytes) -> Dict[str, Any]:
        """Extract text from Word document bytes."""
        result = {
            "content": "",
            "provenance": []
        }
        
        try:
            from io import BytesIO
            doc = docx.Document(BytesIO(docx_bytes))
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    result["content"] += text + "\n"
                    
                    # Create provenance record for paragraph
                    provenance = ProvenanceRecord(
                        page=1,
                        bbox=[0, 0, 0, 0],
                        tool="python-docx",
                        confidence=1.0,
                        element_hash=calculate_element_hash(text, [0, 0, 0, 0], 1),
                        element_type="paragraph",
                        content_preview=text[:50]
                    )
                    result["provenance"].append(provenance)
            
            # Extract text from tables
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    table_content = "\n".join(table_text)
                    result["content"] += f"\n[Table {table_num + 1}]\n{table_content}\n"
                    
                    # Create provenance record for table
                    provenance = ProvenanceRecord(
                        page=1,
                        bbox=[0, 0, 0, 0],
                        tool="python-docx",
                        confidence=1.0,
                        element_hash=calculate_element_hash(table_content, [0, 0, 0, 0], 1),
                        element_type="table",
                        content_preview=table_content[:50]
                    )
                    result["provenance"].append(provenance)
            
        except Exception as e:
            logger.error(f"Word text extraction from bytes failed: {e}")
            raise
        
        return result


class WordTableExtractor:
    """Extract tables from Word documents."""
    
    def __init__(self):
        self.version = "1.0.0"
    
    def get_version(self) -> str:
        return self.version
    
    def extract_tables(self, docx_path: Path) -> Dict[str, Any]:
        """Extract tables from Word document."""
        result = {
            "tables": [],
            "provenance": []
        }
        
        try:
            doc = docx.Document(str(docx_path))
            
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    result["tables"].append({
                        "table_id": f"table_{table_num + 1}",
                        "data": table_data,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0
                    })
                    
                    # Create provenance record for table
                    table_content = "\n".join([" | ".join(row) for row in table_data])
                    provenance = ProvenanceRecord(
                        page=1,
                        bbox=[0, 0, 0, 0],
                        tool="python-docx",
                        confidence=1.0,
                        element_hash=calculate_element_hash(table_content, [0, 0, 0, 0], 1),
                        element_type="table",
                        content_preview=table_content[:50]
                    )
                    result["provenance"].append(provenance)
            
        except Exception as e:
            logger.error(f"Word table extraction failed: {e}")
            raise
        
        return result


class WordImageExtractor:
    """Extract images from Word documents."""
    
    def __init__(self):
        self.version = "1.0.0"
    
    def get_version(self) -> str:
        return self.version
    
    def extract_images(self, docx_path: Path, output_dir: Path) -> Dict[str, Any]:
        """Extract images from Word document."""
        result = {
            "images": [],
            "provenance": []
        }
        
        try:
            # Create images directory
            images_dir = output_dir / "assets"
            images_dir.mkdir(parents=True, exist_ok=True)
            
            # Extract images using docx2txt
            image_list = docx2txt.process(str(docx_path), images_dir)
            
            # Process extracted images
            for img_path in images_dir.glob("*.png"):
                img_info = {
                    "filename": img_path.name,
                    "path": str(img_path),
                    "size": img_path.stat().st_size
                }
                result["images"].append(img_info)
                
                # Create provenance record for image
                provenance = ProvenanceRecord(
                    page=1,
                    bbox=[0, 0, 0, 0],
                    tool="docx2txt",
                    confidence=1.0,
                    element_hash=calculate_element_hash(img_path.name, [0, 0, 0, 0], 1),
                    element_type="image",
                    content_preview=f"Image: {img_path.name}"
                )
                result["provenance"].append(provenance)
            
        except Exception as e:
            logger.error(f"Word image extraction failed: {e}")
            raise
        
        return result


def detect_word_document_characteristics(docx_path: Path) -> Dict[str, Any]:
    """Analyze Word document to detect characteristics."""
    characteristics = {
        "has_tables": False,
        "has_images": False,
        "paragraph_count": 0,
        "table_count": 0,
        "image_count": 0,
        "estimated_domain": None
    }
    
    try:
        doc = docx.Document(str(docx_path))
        
        # Count paragraphs
        characteristics["paragraph_count"] = len([p for p in doc.paragraphs if p.text.strip()])
        
        # Count tables
        characteristics["table_count"] = len(doc.tables)
        characteristics["has_tables"] = characteristics["table_count"] > 0
        
        # Check for images (basic check)
        try:
            from docx2txt import process
            temp_dir = Path("/tmp/word_temp")
            temp_dir.mkdir(exist_ok=True)
            process(str(docx_path), temp_dir)
            image_files = list(temp_dir.glob("*.png"))
            characteristics["image_count"] = len(image_files)
            characteristics["has_images"] = characteristics["image_count"] > 0
            
            # Cleanup
            for img_file in image_files:
                img_file.unlink()
            temp_dir.rmdir()
        except Exception:
            # If image extraction fails, assume no images
            pass
        
        # Basic domain detection based on content
        all_text = " ".join([p.text for p in doc.paragraphs])
        if any(word in all_text.lower() for word in ["research", "study", "analysis", "methodology"]):
            characteristics["estimated_domain"] = "academic"
        elif any(word in all_text.lower() for word in ["business", "report", "financial", "quarterly"]):
            characteristics["estimated_domain"] = "business"
        elif any(word in all_text.lower() for word in ["technical", "specification", "requirements"]):
            characteristics["estimated_domain"] = "technical"
        
    except Exception as e:
        logger.error(f"Word document analysis failed: {e}")
        raise
    
    return characteristics
