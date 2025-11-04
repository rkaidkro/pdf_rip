#!/usr/bin/env python3
"""
Create a test Word document for testing the document processing pipeline.
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path

def create_test_word_document():
    """Create a test Word document with various content types."""
    
    # Create a new document
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document for Document Rip', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a test document to verify Word document processing capabilities.')
    
    doc.add_heading('Text Content', level=1)
    doc.add_paragraph('This section contains regular text content with multiple paragraphs.')
    doc.add_paragraph('The document processing pipeline should be able to extract this text and convert it to Markdown format.')
    
    # Add a table
    doc.add_heading('Sample Table', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add data to the table
    data = [
        ['Name', 'Age', 'City'],
        ['John Doe', '30', 'New York'],
        ['Jane Smith', '25', 'Los Angeles']
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
    
    # Add more content
    doc.add_heading('Lists', level=1)
    doc.add_paragraph('This document contains various types of content:')
    
    # Add a bulleted list
    doc.add_paragraph('• Text paragraphs', style='List Bullet')
    doc.add_paragraph('• Tables with data', style='List Bullet')
    doc.add_paragraph('• Headings and subheadings', style='List Bullet')
    doc.add_paragraph('• Lists (like this one)', style='List Bullet')
    
    # Add a numbered list
    doc.add_paragraph('1. First item in numbered list', style='List Number')
    doc.add_paragraph('2. Second item in numbered list', style='List Number')
    doc.add_paragraph('3. Third item in numbered list', style='List Number')
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('This test document demonstrates the ability to process Word documents with various content types including text, tables, and structured formatting.')
    
    # Save the document
    output_path = Path('examples/test_document.docx')
    doc.save(str(output_path))
    
    print(f"Test Word document created: {output_path}")
    return output_path

if __name__ == "__main__":
    create_test_word_document()
